from __future__ import annotations

from pathlib import Path


def _qn(tag: str) -> str:
    from docx.oxml.ns import qn

    return qn(tag)


def _set_rfonts(style, east_asia: str, latin: str = "Times New Roman") -> None:
    from docx.oxml import OxmlElement

    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(_qn("w:ascii"), latin)
    r_fonts.set(_qn("w:hAnsi"), latin)
    r_fonts.set(_qn("w:eastAsia"), east_asia)
    r_fonts.set(_qn("w:cs"), latin)


def create_template(path: Path) -> Path:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Cm, Pt

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    styles = [
        ("Thesis Heading 1", "黑体", 15, True),
        ("Thesis Heading 2", "黑体", 14, True),
        ("Thesis Heading 3", "黑体", 12, True),
        ("Thesis Body", "宋体", 12, False),
        ("Thesis Abstract", "宋体", 12, False),
        ("Thesis Keywords", "宋体", 12, False),
        ("Thesis Figure Caption", "宋体", 10.5, False),
        ("Thesis Table Caption", "宋体", 10.5, False),
        ("Thesis Reference", "宋体", 10.5, False),
    ]
    for style_name, zh_font, size, bold in styles:
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.size = Pt(size)
        style.font.bold = bold
        _set_rfonts(style, zh_font)
        doc.add_paragraph(f"Sample {style_name}", style=style_name)
    doc.save(path)
    return path


def create_content(path: Path) -> Path:
    from docx import Document

    doc = Document()
    doc.add_paragraph("摘要")
    doc.add_paragraph("本文用于演示论文格式助手的最小工作流。")
    doc.add_paragraph("关键词：论文；格式；Word")
    doc.add_paragraph("1 绪论")
    doc.add_paragraph("这是正文第一段。")
    doc.add_paragraph("1.1 研究背景")
    doc.add_paragraph("这是正文第二段，保留字符级格式。").runs[0].bold = True
    doc.add_paragraph("图1-1 系统流程图")
    doc.add_paragraph("表1-1 示例表")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    doc.add_paragraph("参考文献")
    doc.add_paragraph("[1] Zhang S. Example reference.")
    doc.save(path)
    return path


def generate_examples(root: Path | str = "examples") -> tuple[Path, Path]:
    root = Path(root)
    root.mkdir(parents=True, exist_ok=True)
    template = create_template(root / "template_basic.docx")
    content = create_content(root / "content_basic.docx")
    return template, content


def main() -> int:
    template, content = generate_examples()
    print(f"Generated {template}")
    print(f"Generated {content}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

