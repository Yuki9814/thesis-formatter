from __future__ import annotations

import base64
import shutil
import zipfile
from pathlib import Path

import pytest


PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
)


def _docx_qn(tag: str) -> str:
    from docx.oxml.ns import qn

    return qn(tag)


def _set_style_rfonts(style, east_asia: str, latin: str = "Times New Roman") -> None:
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        from docx.oxml import OxmlElement

        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(_docx_qn("w:ascii"), latin)
    r_fonts.set(_docx_qn("w:hAnsi"), latin)
    r_fonts.set(_docx_qn("w:eastAsia"), east_asia)
    r_fonts.set(_docx_qn("w:cs"), latin)


def _make_template(path: Path) -> Path:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Cm, Pt

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    specs = [
        ("Thesis Heading 1", "黑体", 15, True),
        ("Thesis Heading 2", "黑体", 14, True),
        ("Thesis Heading 3", "黑体", 12, True),
        ("Thesis Body", "宋体", 12, False),
        ("Thesis Abstract", "宋体", 12, False),
        ("Thesis Keywords", "宋体", 12, False),
        ("Thesis Figure Caption", "宋体", 10.5, False),
        ("Thesis Table Caption", "宋体", 10.5, False),
        ("Thesis Reference", "宋体", 10.5, False),
    ]
    for name, font, size, bold in specs:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.size = Pt(size)
        style.font.bold = bold
        _set_style_rfonts(style, font)
        p = doc.add_paragraph(f"Sample {name}", style=name)
        if "Heading" in name:
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
    doc.save(path)
    return path


def _make_simple_content(path: Path, image: bool = False, table: bool = False) -> Path:
    from docx import Document

    doc = Document()
    doc.add_paragraph("摘要")
    doc.add_paragraph("这是一段摘要内容。")
    doc.add_paragraph("关键词：格式；论文；Word")
    doc.add_paragraph("1 绪论")
    doc.add_paragraph("正文第一段，包含普通文字。")
    doc.add_paragraph("1.1 研究背景")
    doc.add_paragraph("第二段正文，包含加粗文字。").runs[0].bold = True
    doc.add_paragraph("图1-1 系统流程图")
    doc.add_paragraph("表1-1 测试表")
    if table:
        tbl = doc.add_table(rows=2, cols=2)
        tbl.cell(0, 0).text = "A"
        tbl.cell(0, 1).text = "B"
    if image:
        img = path.with_suffix(".png")
        img.write_bytes(PNG_1X1)
        doc.add_picture(str(img))
    doc.add_paragraph("参考文献")
    doc.add_paragraph("[1] Zhang S. A reference item.")
    doc.save(path)
    return path


def _rewrite_docx(path: Path, replacements: dict[str, bytes], additions: dict[str, bytes] | None = None) -> None:
    tmp = path.with_suffix(".rewrite.docx")
    with zipfile.ZipFile(path, "r") as src, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as dst:
        for item in src.infolist():
            data = replacements.get(item.filename, src.read(item.filename))
            dst.writestr(item, data)
        for name, data in (additions or {}).items():
            if name not in src.namelist():
                dst.writestr(name, data)
    shutil.move(tmp, path)


def _make_advanced_content(path: Path) -> Path:
    from lxml import etree

    from core.xml_utils import NS, parse_xml, qn, to_xml

    _make_simple_content(path, image=True, table=True)
    with zipfile.ZipFile(path) as package:
        root = parse_xml(package.read("word/document.xml"))
    paragraphs = root.xpath(".//w:body//w:p", namespaces=NS)
    target = paragraphs[1]
    bookmark_start = etree.Element(qn("w:bookmarkStart"))
    bookmark_start.set(qn("w:id"), "1")
    bookmark_start.set(qn("w:name"), "bm_test")
    bookmark_end = etree.Element(qn("w:bookmarkEnd"))
    bookmark_end.set(qn("w:id"), "1")
    target.insert(0, bookmark_start)
    target.append(bookmark_end)

    toc_p = etree.Element(qn("w:p"))
    run = etree.SubElement(toc_p, qn("w:r"))
    instr = etree.SubElement(run, qn("w:instrText"))
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    root.find("./w:body", namespaces=NS).insert(0, toc_p)

    math_p = etree.Element(qn("w:p"))
    math_run = etree.SubElement(math_p, qn("w:r"))
    omath = etree.SubElement(math_run, f"{{{NS['m']}}}oMath")
    mr = etree.SubElement(omath, f"{{{NS['m']}}}r")
    mt = etree.SubElement(mr, f"{{{NS['m']}}}t")
    mt.text = "x=1"
    root.find("./w:body", namespaces=NS).append(math_p)

    _rewrite_docx(
        path,
        {"word/document.xml": to_xml(root)},
        {
            "word/footnotes.xml": b'<?xml version="1.0" encoding="UTF-8"?><w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>',
            "word/endnotes.xml": b'<?xml version="1.0" encoding="UTF-8"?><w:endnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>',
        },
    )
    return path


def _make_bad_reference(path: Path) -> Path:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    for idx in range(12):
        p = doc.add_paragraph(f"手动格式段落 {idx}")
        p.paragraph_format.first_line_indent = Pt(24)
        p.paragraph_format.space_after = Pt(6)
        run = p.runs[0]
        run.bold = idx % 2 == 0
        run.font.size = Pt(12 + idx % 3)
    doc.save(path)
    return path


@pytest.fixture()
def simple_template(tmp_path: Path) -> Path:
    return _make_template(tmp_path / "template.docx")


@pytest.fixture()
def simple_content(tmp_path: Path) -> Path:
    return _make_simple_content(tmp_path / "content.docx")


@pytest.fixture()
def image_content(tmp_path: Path) -> Path:
    return _make_simple_content(tmp_path / "content_image.docx", image=True)


@pytest.fixture()
def table_content(tmp_path: Path) -> Path:
    return _make_simple_content(tmp_path / "content_table.docx", table=True)


@pytest.fixture()
def advanced_content(tmp_path: Path) -> Path:
    return _make_advanced_content(tmp_path / "content_advanced.docx")


@pytest.fixture()
def bad_reference(tmp_path: Path) -> Path:
    return _make_bad_reference(tmp_path / "bad_reference.docx")

